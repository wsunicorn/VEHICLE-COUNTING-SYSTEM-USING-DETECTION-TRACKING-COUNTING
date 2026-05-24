from pathlib import Path
import json

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "docs" / "Nhom19_paper_ComputerVision.docx"
OUT_DIR = ROOT / "dtc_counting" / "outputs" / "final_cam5_b1_b4_20260524_v2"


def element_text(el):
    return "".join((t.text or "") for t in el.xpath(".//w:t")).strip()


def remove_generated_content(doc):
    body = doc.element.body
    children = list(body)

    def is_update(txt):
        return txt.startswith("CẬP NHẬT THỰC NGHIỆM NGÀY 24/05/2026") or txt.startswith(
            "C?P NH?T TH?C NGHI?M NG?Y 24/05/2026"
        )

    def is_refs(txt):
        return txt.startswith("TÀI LIỆU THAM KHẢO") or txt.startswith("T?I LI?U THAM KH?O")

    starts = [i for i, el in enumerate(children) if is_update(element_text(el))]
    for start in reversed(starts):
        end = None
        for idx in range(start + 1, len(children)):
            if is_refs(element_text(children[idx])):
                end = idx
                break
        if end is None:
            end = len(children) - 1 if children and children[-1].tag.endswith("sectPr") else len(children)
        for el in children[start:end]:
            if el in body:
                body.remove(el)

    for el in list(body):
        txt = element_text(el)
        if txt.startswith("[11] M. Kocur") or txt.startswith("[12] D. Gloudemans"):
            body.remove(el)


def find_refs_element(doc):
    body = doc.element.body
    for el in list(body):
        txt = element_text(el)
        if txt.startswith("TÀI LIỆU THAM KHẢO") or txt.startswith("T?I LI?U THAM KH?O"):
            return el
    return list(body)[-1]


def update_existing_result_table(doc):
    for table in doc.tables:
        for row in table.rows:
            first = row.cells[0].text.strip() if row.cells else ""
            if first.startswith("B3: SAM Auto"):
                row.cells[0].text = "B3: SAM Auto (quality gate chưa đạt)"
            if first.startswith("B4: Grounded SAM") or first.startswith("B4: Grounding DINO + SAM"):
                values = [
                    "B4: Grounding DINO + SAM (tự động, cần hậu kiểm)",
                    "125",
                    "0.4071",
                    "0.4150*",
                    "69.79%",
                    "5.11",
                ]
                for cell, val in zip(row.cells, values):
                    cell.text = val


def add_update_section(doc, refs_el):
    summary = json.loads((OUT_DIR / "comparison_summary.json").read_text(encoding="utf-8"))
    b3 = json.loads((OUT_DIR / "b3_sam_auto_bootstrap.json").read_text(encoding="utf-8"))
    b4 = json.loads((OUT_DIR / "b4_grounded_sam_bootstrap.json").read_text(encoding="utf-8"))
    styles = {s.name for s in doc.styles}

    def style(name):
        return name if name in styles else None

    def before_refs(el):
        refs_el.addprevious(el)

    def p(text="", style_name=None, bold=False, italic=False, center=False):
        para = doc.add_paragraph(style=style(style_name) if style_name else None)
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        before_refs(para._p)
        return para

    def heading(text, level=1):
        p(text, f"Heading {level}")

    def table(rows, cols, caption, headers):
        p(caption, italic=True, center=True)
        tbl = doc.add_table(rows=rows, cols=cols)
        if "Table Grid" in styles:
            tbl.style = "Table Grid"
        for idx, header in enumerate(headers):
            tbl.rows[0].cells[idx].text = header
            for para in tbl.rows[0].cells[idx].paragraphs:
                for run in para.runs:
                    run.bold = True
        before_refs(tbl._tbl)
        return tbl

    def picture(path, caption, width=5.9):
        path = Path(path)
        if not path.exists():
            return
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(str(path), width=Inches(width))
        before_refs(para._p)
        p(caption, italic=True, center=True)

    def pct(x):
        return f"{x * 100:.2f}%"

    heading("CẬP NHẬT THỰC NGHIỆM NGÀY 24/05/2026", 1)
    p(
        "Hệ thống được xây dựng như một pipeline tổng quát cho bài toán đếm xe đa lớp, "
        "đa hướng trên dữ liệu camera giao thông của AI City Challenge. Các mô-đun chính "
        "bao gồm cấu hình ROI/MOI, phát hiện xe bằng YOLO, theo dõi đa đối tượng, gán "
        "MOI theo quỹ đạo và đánh giá bằng các chỉ số nwRMSE, S1 Effectiveness, Count "
        "Accuracy và MAE. Kết quả định lượng trong báo cáo dùng mẫu đánh giá local có "
        "ground truth để kiểm chứng pipeline, trong khi kiến trúc và web demo được thiết "
        "kế để áp dụng cho các camera khác khi có ROI/MOI hoặc cấu hình bootstrap tương ứng."
    )
    p(
        "Cách trình bày dưới đây nhấn mạnh tính tổng quát của phương pháp, nhưng vẫn giữ "
        "ranh giới rõ giữa kết quả định lượng đáng tin cậy và các nhánh tự động cần hậu "
        "kiểm trực quan."
    )

    heading("Sơ đồ hệ thống", 2)
    picture(
        ROOT / "docs" / "workflow_diagram.png",
        "Hình 1. Sơ đồ tổng quan mới của hệ thống: cấu hình ROI/MOI, YOLO, tracking, gán MOI, xuất kết quả và đánh giá B1-B4.",
        width=6.2,
    )

    heading("Vai trò của SAM trong hệ thống", 2)
    p(
        "SAM được dùng ở tầng bootstrap ROI/MOI, không thay thế detector đếm xe. Trong "
        "SAM Automatic, hệ thống lấy frame đại diện, sinh mask ứng viên, chấm điểm vùng "
        "mặt đường, loại các vùng có đặc trưng cây cỏ và dựng ROI polygon. Với nhánh "
        "Grounding DINO + SAM, Grounding DINO tạo vùng gợi ý bằng prompt ngôn ngữ như "
        "road surface, traffic lane, intersection; SAM sau đó tinh chỉnh mask để tạo ROI. "
        "Các MOI vector được tạo từ quỹ đạo track hoặc từ hình học mask, rồi đi qua quality "
        "gate: kiểm tra kích thước ROI, số vector MOI hợp lệ và nguy cơ fallback toàn khung."
    )
    p(
        "Do ROI/MOI quyết định trực tiếp việc gán movement, các nhánh SAM hiện được xem là "
        "phương pháp hỗ trợ khởi tạo nhanh. Với báo cáo và demo ổn định, cấu hình thủ công "
        "hoặc file ROI/MOI chuẩn vẫn là baseline tin cậy; SAM/Grounded-SAM được đưa vào như "
        "nhánh tự động hóa có điều kiện hậu kiểm."
    )

    headers = ["Nhánh", "Cơ chế", "Ưu điểm", "Giới hạn hiện tại", "Cách dùng trong báo cáo"]
    tbl = table(4, len(headers), "Bảng 1. Tổng hợp vai trò của SAM và Grounded-SAM trong bootstrap ROI/MOI.", headers)
    sam_rows = [
        [
            "SAM Automatic",
            "Sinh mask tự động từ frame đại diện, chọn mask mặt đường và dựng ROI/MOI.",
            "Không cần prompt, có thể tạo cấu hình ban đầu nhanh.",
            f"Quality gate báo {b3['quality']['status']}; chỉ có {b3['quality']['valid_moi_count']} MOI hợp lệ trong lần chạy local.",
            "Dùng minh họa hướng tự động hóa, không dùng làm kết quả định lượng chính nếu thiếu MOI.",
        ],
        [
            "Grounding DINO + SAM",
            "Grounding DINO định vị vùng đường bằng prompt, SAM phân đoạn chi tiết vùng ROI.",
            "Có hướng dẫn ngữ nghĩa, phù hợp hơn khi frame phức tạp.",
            f"Lần chạy hiện tại sinh {b4['quality']['valid_moi_count']} MOI từ bootstrap; pipeline dùng track-mined MOI fallback để tránh gán hướng bằng vector thiếu.",
            "Dùng như baseline tự động có hậu xử lý, cần kiểm tra ROI/MOI trước khi demo.",
        ],
        [
            "Quality Gate",
            "Kiểm tra ROI quá rộng, fallback toàn khung, vùng cây cỏ và số MOI tối thiểu.",
            "Giảm khả năng đưa cấu hình sai vào bước đếm.",
            "Chưa thay được kiểm tra trực quan của người dùng.",
            "Giải thích vì sao B3/B4 có thể N/A hoặc điểm thấp.",
        ],
    ]
    for ridx, row in enumerate(sam_rows, start=1):
        for cidx, value in enumerate(row):
            tbl.rows[ridx].cells[cidx].text = value

    heading("Kết quả thực nghiệm B1-B4", 2)
    headers = ["Baseline", "Mô tả", "Pred/GT", "nwRMSE", "S1 Eff.", "S1 Overall*", "Accuracy", "MAE", "Kết luận"]
    tbl = table(5, len(headers), "Bảng 2. Kết quả so sánh các cấu hình baseline trên mẫu đánh giá local.", headers)
    rows = [
        [
            "B1",
            summary["b1"]["label"],
            f"{summary['b1']['pred_total']}/{summary['b1']['gt_total']}",
            f"{summary['b1']['nwRMSE']:.4f}",
            f"{summary['b1']['S1_Effectiveness']:.4f}",
            f"{summary['b1']['S1_Overall']:.4f}",
            pct(summary["b1"]["count_accuracy"]),
            f"{summary['b1']['mae']:.2f}",
            "Baseline định lượng chính, ROI/MOI được kiểm chứng.",
        ],
        [
            "B2",
            summary["b2"]["label"],
            f"{summary['b2']['pred_total']}/{summary['b2']['gt_total']}",
            f"{summary['b2']['nwRMSE']:.4f}",
            f"{summary['b2']['S1_Effectiveness']:.4f}",
            f"{summary['b2']['S1_Overall']:.4f}",
            pct(summary["b2"]["count_accuracy"]),
            f"{summary['b2']['mae']:.2f}",
            "Hợp lệ nhưng phân bổ theo movement kém ổn định hơn B1.",
        ],
        [
            "B3",
            "SAM Automatic",
            "N/A",
            "N/A",
            "N/A",
            "N/A",
            "N/A",
            "N/A",
            "Không đưa vào định lượng khi quality gate chưa đạt.",
        ],
        [
            "B4",
            summary["b4"]["label"],
            f"{summary['b4']['pred_total']}/{summary['b4']['gt_total']}",
            f"{summary['b4']['nwRMSE']:.4f}",
            f"{summary['b4']['S1_Effectiveness']:.4f}",
            f"{summary['b4']['S1_Overall']:.4f}",
            pct(summary["b4"]["count_accuracy"]),
            f"{summary['b4']['mae']:.2f}",
            "Cải thiện sau khi dùng track-mined MOI fallback cho ROI bootstrap.",
        ],
    ]
    for ridx, row in enumerate(rows, start=1):
        for cidx, value in enumerate(row):
            tbl.rows[ridx].cells[cidx].text = value
    p(
        "*S1 Overall là giá trị xấp xỉ nội bộ khi chưa dùng bộ chấm leaderboard chính thức. "
        "Trong thảo luận, nwRMSE và S1 Effectiveness được dùng làm chỉ báo chính cho chất lượng đếm."
    )

    heading("Minh họa kết quả", 2)
    picture(
        ROOT / "dtc_counting" / "web_demo" / "media" / "20260524_155110" / "bootstrap_overlay.jpg",
        "Hình 2. Minh họa nhánh SAM trajectory fallback: tạo ROI/MOI tự động và cần người dùng kiểm tra trước khi đếm.",
        width=5.8,
    )
    picture(
        OUT_DIR / "b4_grounded_sam_overlay.jpg",
        "Hình 3. Minh họa nhánh Grounding DINO + SAM: ROI được bootstrap tự động nhưng MOI cần được hậu xử lý để ổn định hơn.",
        width=5.8,
    )

    heading("So sánh với nghiên cứu liên quan", 2)
    headers = ["Phương pháp", "Phạm vi", "Kết quả công bố", "Nhận xét cho hệ thống hiện tại"]
    tbl = table(5, len(headers), "Bảng 3. So sánh định hướng với các nghiên cứu trên AI City Challenge Track 1.", headers)
    rows = [
        [
            "Tiny-PIRATE [1]",
            "AI City Challenge 2021 Track 1",
            "S1=0.9459, hạng 2 theo báo cáo tác giả.",
            "Cho thấy pipeline tối ưu theo benchmark và MOI ổn định có ảnh hưởng lớn đến điểm cuối.",
        ],
        [
            "CenterTrack-based counting [11]",
            "AI City Challenge 2021 Track 1",
            "S1=0.8449, hạng 8 public leaderboard.",
            "Nhấn mạnh vai trò của tracker mạnh và gán hướng đáng tin cậy.",
        ],
        [
            "LBT-Count [12]",
            "AI City Challenge 2021 Track 1",
            "Báo cáo nhanh hơn khoảng 52% so với tracker baseline và đạt hạng 7 trên public set.",
            "Gợi ý hướng tối ưu tốc độ bằng localization/tracking nhẹ hơn.",
        ],
        [
            "Hệ thống đề xuất",
            "Pipeline local có thể cấu hình cho nhiều camera; bảng số liệu minh họa trên mẫu có ground truth.",
            f"B1 đạt nwRMSE={summary['b1']['nwRMSE']:.4f}, S1_Eff.={summary['b1']['S1_Effectiveness']:.4f}.",
            "Phù hợp làm hệ thống demo và nền tảng mở rộng; nhánh SAM giúp giảm công cấu hình khi kết hợp quality gate/fallback.",
        ],
    ]
    for ridx, row in enumerate(rows, start=1):
        for cidx, value in enumerate(row):
            tbl.rows[ridx].cells[cidx].text = value

    heading("Kết luận cập nhật", 2)
    p(
        "Kết quả hiện tại cho thấy pipeline phát hiện-theo dõi-đếm đã hoạt động hợp lý, "
        "đặc biệt ở nhánh B1 với ROI/MOI đã kiểm chứng. B2 cho thấy khả năng khai thác "
        "MOI từ quỹ đạo, còn B3/B4 mở ra hướng tự động hóa cấu hình bằng SAM nhưng cần "
        "quality gate và hậu kiểm. Vì vậy, báo cáo nên trình bày hệ thống như một kiến "
        "trúc tổng quát có nhiều chế độ cấu hình ROI/MOI, trong đó B1/B2 là kết quả định "
        "lượng chính và B3/B4 là phần mở rộng thử nghiệm có phân tích hạn chế rõ ràng."
    )


def add_references(doc, refs_el):
    body = doc.element.body
    link_el = None
    for el in list(body):
        if element_text(el).startswith("Link github"):
            link_el = el
            break
    if link_el is None:
        link_el = refs_el

    refs = [
        '[11] M. Kocur, T. Dwornik, and M. Koziarski, "Multi-Class Multi-Movement Vehicle Counting Based on CenterTrack," CVPR Workshops AI City Challenge, 2021.',
        '[12] D. Gloudemans et al., "Fast Vehicle Turning-Movement Counting Using Localization-Based Tracking," CVPR Workshops AI City Challenge, 2021.',
    ]
    for ref in refs:
        para = doc.add_paragraph(ref)
        link_el.addprevious(para._p)


def main():
    doc = Document(str(DOC_PATH))
    remove_generated_content(doc)
    refs_el = find_refs_element(doc)
    update_existing_result_table(doc)
    add_update_section(doc, refs_el)
    add_references(doc, refs_el)
    doc.save(str(DOC_PATH))
    print(f"Updated {DOC_PATH}")


if __name__ == "__main__":
    main()
